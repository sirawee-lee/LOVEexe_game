#!/usr/bin/env python3
"""Generate two .docx files for LOVE.EXE Group 23."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def table_header_row(table, headers, bg='2E4057'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_table_row(table, values, bg=None):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        if bg:
            set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
    return row

def hline(doc):
    doc.add_paragraph('─' * 80)

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 1 — PMP + SRS
# ══════════════════════════════════════════════════════════════════════════════

def make_pmp_srs():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Cover ─────────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run('LOVE.EXE')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('Project Management Plan  &  Software Requirements Specification')
    r2.font.size = Pt(13)
    r2.italic = True

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run('Software Studio 114  |  Group 23  |  Version 1.0  |  May 2026')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ── 1. Team & Roles ───────────────────────────────────────────────────────
    heading(doc, '1. Team & Role Assignments', level=1)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    table_header_row(tbl, ['Student ID', 'Name', 'Role', 'Responsibilities'])

    members = [
        ('110006226', 'Sirawee',   'Lead Developer\n(Framework)',
         'GameManager, PlayerController, BoothTrigger,\nDialogueSystem, EndingManager, Git repo management,\nScene integration & final merge'),
        ('111062119', '蘇昱銓',    'Mini-game Developer A',
         'MiniGame_Professor (Logic Gate Puzzle)\nMiniGame_NiuPai (Grid Stealth / AI Chase)'),
        ('113006230', 'Anan',      'Mini-game Developer B',
         'MiniGame_Father (Rhythm / Catch)\nEnding scene layout & credits roll'),
        ('109006229', '張淑芬',    'Pixel Artist — World',
         'NTHU campus map tileset, building sprites,\nenvironment tiles (trees, pond, track field)'),
        ('109006258', '許夏綠',    'Pixel Artist — Characters & UI',
         'Character sprites (Boy, Mei, Professor, Father, Niu Pai),\nHUD icons (hearts, coins), UI frames, ending CG'),
    ]

    row_colors = ['FFFFFF', 'F2F2F2']
    for i, m in enumerate(members):
        add_table_row(tbl, m, bg=row_colors[i % 2])

    doc.add_paragraph()

    # ── 2. Project Overview ───────────────────────────────────────────────────
    heading(doc, '2. Project Overview', level=1)

    info_tbl = doc.add_table(rows=5, cols=2)
    info_tbl.style = 'Table Grid'
    info_data = [
        ('Game Title', 'LOVE.EXE'),
        ('Engine', 'Cocos Creator 2.4.8  (JavaScript — mandatory)'),
        ('Platform', 'Web Browser'),
        ('Deadline', '30 May 2026'),
        ('Repository', 'GitHub — ask Sirawee for invite link'),
    ]
    for i, (k, v) in enumerate(info_data):
        row = info_tbl.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        set_cell_bg(row.cells[0], 'D6EAF8')
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    body(doc,
         'LOVE.EXE is an 8–16 bit pixel-art RPG set on NTHU campus. '
         'The player (a nerdy EECS boy) explores a top-down campus map, '
         'challenges NPCs in three mini-games, collects Heart Fragments, '
         'and reaches one of five distinct endings depending on performance and dialogue choices.',
         size=10)

    # ── 3. Scope ──────────────────────────────────────────────────────────────
    heading(doc, '3. Scope', level=1)
    heading(doc, '3.1  In-Scope (Must Ship)', level=2)
    in_scope = [
        'Top-down NTHU campus map with walkable paths and 3 booth entry points',
        'WASD movement + E/Space interaction for player character',
        'Three fully playable mini-games (Professor, Niu Pai, Father)',
        'Heart Fragment system (0–3) + Coin reward system',
        'Affinity score tracked via dialogue choices (0–100)',
        'Five endings: True Love, Friend Zone, Rejected, Niu Pai Easter Egg, EECS Overload Easter Egg',
        'Data-driven dialogue loaded from game-data.json (no hard-coded strings)',
        'HUD showing hearts and coins at all times',
        'Seamless scene transitions (map ↔ mini-game ↔ ending)',
    ]
    for s in in_scope:
        bullet(doc, s)

    heading(doc, '3.2  Out-of-Scope (Do NOT implement now)', level=2)
    out_scope = [
        'Online multiplayer / global chat',
        'Multiple romance routes (single love interest only)',
        'Leaderboard backend (local high-score display is OK)',
        'Mobile / touch controls',
    ]
    for s in out_scope:
        bullet(doc, s)

    # ── 4. Functional Requirements ────────────────────────────────────────────
    heading(doc, '4. Functional Requirements', level=1)

    req_tbl = doc.add_table(rows=1, cols=4)
    req_tbl.style = 'Table Grid'
    table_header_row(req_tbl, ['ID', 'Requirement', 'Owner', 'Priority'])

    reqs = [
        ('FR-01', 'Player can move in 4 directions via WASD. Movement speed is 150 px/s.',                  'Sirawee',  'High'),
        ('FR-02', 'Player approaching a booth sees a "Press E" prompt and can enter the mini-game.',         'Sirawee',  'High'),
        ('FR-03', 'GameManager persists across scenes (PersistRootNode). Hearts/coins never reset mid-run.','Sirawee',  'High'),
        ('FR-04', 'Mini-game Professor: solve ≥3 logic-gate circuits before 60-second timer expires.',       '蘇昱銓',   'High'),
        ('FR-05', 'Mini-game Niu Pai: escape the AI-chasing dog on a grid map (reach exit without caught).', '蘇昱銓',   'High'),
        ('FR-06', 'Mini-game Father: catch good items / dodge bad items to reach 60 pts in 45 seconds.',     'Anan',     'High'),
        ('FR-07', 'Completing any mini-game for the first time awards 1 Heart Fragment.',                     'Sirawee',  'High'),
        ('FR-08', 'Dialogue boxes support sequential lines, typewriter effect, and 2-choice branches.',       'Sirawee',  'High'),
        ('FR-09', 'Each dialogue choice changes affinityScore by a defined delta (stored in JSON).',          'Sirawee',  'Medium'),
        ('FR-10', 'At game end, resolveEnding() selects 1 of 5 endings based on hearts + affinityScore.',    'Sirawee',  'High'),
        ('FR-11', 'Ending scene shows slides read from game-data.json (title, body text, CG image).',        'Anan',     'Medium'),
        ('FR-12', 'HUD updates heart icons and coin count in real-time via cc.systemEvent.',                  'Sirawee',  'High'),
        ('FR-13', 'Campus map tiles drawn as 16×16 or 32×32 px sprites in a unified tileset.',               '張淑芬',   'High'),
        ('FR-14', 'All character sprites include idle and walk animations for 4 directions.',                 '許夏綠',   'High'),
    ]

    row_colors = ['FFFFFF', 'F2F2F2']
    priority_colors = {'High': 'FADBD8', 'Medium': 'FEF9E7', 'Low': 'EAFAF1'}
    for i, r in enumerate(reqs):
        row = add_table_row(req_tbl, r, bg=row_colors[i % 2])
        pri_cell = row.cells[3]
        set_cell_bg(pri_cell, priority_colors.get(r[3], 'FFFFFF'))

    doc.add_paragraph()

    # ── 5. Mini-game Specifications ───────────────────────────────────────────
    heading(doc, '5. Mini-game Specifications', level=1)

    # 5.1 Professor
    heading(doc, '5.1  MiniGame_Professor — "Circuit Pulse"', level=2)
    body(doc, 'Owner: 蘇昱銓  |  Scene: MiniGame_Professor  |  Base class: MiniGameBase.js', bold=True)
    specs = [
        ('Mechanic',    'A partially broken logic-gate circuit is displayed. Player must click/drag to place missing gates (AND/OR/NOT) so the output matches the target value. Multiple rounds with increasing complexity.'),
        ('Win Condition', 'Solve 3 circuits correctly before the 60-second timer runs out.'),
        ('Lose Condition', 'Timer reaches 0 with fewer than 3 circuits solved, OR 3 wrong answers submitted.'),
        ('Score',       '+20 pts per correct circuit. Bonus +10 if solved in < 5 seconds.'),
        ('Heart Award', 'First-time completion → +1 Heart Fragment (handled by MiniGameBase → GameManager).'),
        ('Key Classes', 'MiniGame_Professor.js  |  CircuitNode.js  |  GateEvaluator.js'),
    ]
    stbl = doc.add_table(rows=len(specs), cols=2)
    stbl.style = 'Table Grid'
    for i, (k, v) in enumerate(specs):
        stbl.rows[i].cells[0].text = k
        stbl.rows[i].cells[1].text = v
        set_cell_bg(stbl.rows[i].cells[0], 'D6EAF8')
        for cell in stbl.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    # 5.2 Niu Pai
    heading(doc, '5.2  MiniGame_NiuPai — "Niu Pai\'s Grid Stealth"', level=2)
    body(doc, 'Owner: 蘇昱銓  |  Scene: MiniGame_NiuPai  |  Base class: MiniGameBase.js', bold=True)
    specs2 = [
        ('Mechanic',    'Grid-based map (15×10 tiles). Player moves tile-by-tile with WASD. Niu Pai chases using simplified A* pathfinding. Player can place 3 "sausage" decoys to temporarily redirect the dog.'),
        ('Win Condition', 'Reach the exit tile without being caught within 90 seconds.'),
        ('Lose Condition', 'Niu Pai occupies the same tile as the player, OR timer runs out.'),
        ('Score',       '+50 pts for reaching exit. +5 pts for each remaining second. +10 per decoy unused.'),
        ('Easter Egg Trigger', 'If player feeds Niu Pai (walks into it) 3 times across any visits → sets GameManager.easterEggs.fedNiupai3Times = true.'),
        ('Key Classes', 'MiniGame_NiuPai.js  |  GridMap.js  |  NiuPaiAI.js  |  AStarPathfinder.js'),
    ]
    stbl2 = doc.add_table(rows=len(specs2), cols=2)
    stbl2.style = 'Table Grid'
    for i, (k, v) in enumerate(specs2):
        stbl2.rows[i].cells[0].text = k
        stbl2.rows[i].cells[1].text = v
        set_cell_bg(stbl2.rows[i].cells[0], 'D5F5E3')
        for cell in stbl2.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    # 5.3 Father
    heading(doc, '5.3  MiniGame_Father — "Heartbeat Sync"', level=2)
    body(doc, 'Owner: Anan  |  Scene: MiniGame_Father  |  Base class: MiniGameBase.js', bold=True)
    specs3 = [
        ('Mechanic',    'Items fall from the top of the screen. Player moves a basket left/right (mouse or A/D). Catch GOOD items (flowers, grades, sincerity tokens) and dodge BAD items (complaints, debt bills). Uses Object Pool for performance.'),
        ('Win Condition', 'Reach 60 points before 45 seconds expire.'),
        ('Lose Condition', 'Catch 3 BAD items (lives = 0), or timer runs out below 60 pts.'),
        ('Score',       '+10 pts per GOOD item. -1 life per BAD item. Speed increases every 15 seconds.'),
        ('Heart Award', 'First-time completion → +1 Heart Fragment.'),
        ('Key Classes', 'MiniGame_Father.js  |  ItemSpawner.js  |  ObjectPool.js'),
    ]
    stbl3 = doc.add_table(rows=len(specs3), cols=2)
    stbl3.style = 'Table Grid'
    for i, (k, v) in enumerate(specs3):
        stbl3.rows[i].cells[0].text = k
        stbl3.rows[i].cells[1].text = v
        set_cell_bg(stbl3.rows[i].cells[0], 'FDEBD0')
        for cell in stbl3.rows[i].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    # ── 6. Ending System ──────────────────────────────────────────────────────
    heading(doc, '6. Ending System', level=1)
    body(doc, 'Endings are resolved by GameManager.resolveEnding() at the end of the game. '
              'The result key is passed to EndingManager.js which loads slides from game-data.json.', size=10)
    doc.add_paragraph()

    end_tbl = doc.add_table(rows=1, cols=4)
    end_tbl.style = 'Table Grid'
    table_header_row(end_tbl, ['Ending Key', 'Trigger Condition', 'Type', 'Owner'])

    endings = [
        ('ending_true_love',     'hearts == 3  AND  affinityScore >= 70',              'Normal',    'Sirawee'),
        ('ending_friend_zone',   'hearts >= 2  OR  affinityScore >= 40  (not true love)', 'Normal', 'Sirawee'),
        ('ending_rejected',      'All other cases (low score, low affinity)',           'Normal',    'Sirawee'),
        ('ending_niupai',        'fedNiupai3Times == true  AND  affinityScore < 20',    'Easter Egg', 'Anan'),
        ('ending_eecs_overload', 'neverRomantic == true  AND  allMiniGamesCompleted',   'Easter Egg', 'Anan'),
    ]
    end_colors = {'Normal': 'FFFFFF', 'Easter Egg': 'FFF3CD'}
    for e in endings:
        row = add_table_row(end_tbl, e)
        set_cell_bg(row.cells[2], end_colors.get(e[2], 'FFFFFF'))

    doc.add_paragraph()

    # ── 7. Art Asset Requirements ─────────────────────────────────────────────
    heading(doc, '7. Art Asset Requirements', level=1)

    heading(doc, '7.1  Pixel Art Specs', level=2)
    bullet(doc, 'Base tile size: 32 × 32 px')
    bullet(doc, 'Character sprite sheet: 4 directions × (idle + 2-frame walk) = 12 frames minimum per character')
    bullet(doc, 'Export format: PNG with transparent background')
    bullet(doc, 'Color palette: 16–32 colors per character to keep the 8-bit feel consistent')
    bullet(doc, 'File naming: snake_case — e.g. boy_walk_down.png, girl_idle.png')

    heading(doc, '7.2  Asset Checklist — 張淑芬 (World)', level=2)
    world_assets = [
        'tileset_ground.png — grass, stone path, indoor floor',
        'tileset_objects.png — trees, fountain, pond, bench, gazebo',
        'tileset_buildings.png — Delta building facade, classroom exterior',
        'map_overview.png — full campus overview (used as background reference)',
        'booth_professor.png, booth_niupai.png, booth_father.png — booth signs',
    ]
    for a in world_assets:
        bullet(doc, a)

    heading(doc, '7.3  Asset Checklist — 許夏綠 (Characters & UI)', level=2)
    char_assets = [
        'boy_spritesheet.png — 4-dir walk + idle (protagonist)',
        'girl_spritesheet.png — Mei idle pose (dialogue avatar)',
        'professor_spritesheet.png',
        'father_spritesheet.png',
        'niupai_spritesheet.png — dog walk + run 4-dir',
        'ui_heart_full.png, ui_heart_empty.png',
        'ui_coin.png',
        'dialogue_frame.png — 9-slice compatible panel',
        'btn_choice.png — 9-slice button background',
        'ending_cg_*.png — 2 images per ending = 10 images total',
    ]
    for a in char_assets:
        bullet(doc, a)

    heading(doc, '7.4  Folder placement', level=2)
    bullet(doc, 'All textures go into: assets/textures/characters/  or  assets/textures/map/  or  assets/textures/ui/')
    bullet(doc, 'Ending CGs: assets/textures/endings/')
    bullet(doc, 'DO NOT rename files after Sirawee has wired them in code — coordinate first.')

    # ── 8. Sprint Plan ────────────────────────────────────────────────────────
    heading(doc, '8. Sprint Plan (3 Weeks to May 30)', level=1)

    sp_tbl = doc.add_table(rows=1, cols=4)
    sp_tbl.style = 'Table Grid'
    table_header_row(sp_tbl, ['Sprint', 'Dates', 'Goals', 'Deliverables'])

    sprints = [
        ('Sprint 1\n(Foundation)',
         'May 9–15',
         '• Sirawee: open Cocos project, wire GameManager + PlayerController into MainMap scene\n'
         '• 蘇昱銓 & Anan: study MiniGameBase, design mini-game flowcharts\n'
         '• 張淑芬: deliver ground tileset + 1 building tileset\n'
         '• 許夏綠: deliver boy + girl sprites',
         'Playable character walking on map\nBasic tileset applied\nAll sprites for player/girl done'),

        ('Sprint 2\n(Mini-games)',
         'May 16–24',
         '• 蘇昱銓: implement MiniGame_Professor & MiniGame_NiuPai (prototype logic)\n'
         '• Anan: implement MiniGame_Father + Ending scene layout\n'
         '• Sirawee: integrate all 3 mini-games with GameManager, test transitions\n'
         '• Artists: deliver NPC sprites + UI assets',
         'All 3 mini-games playable end-to-end\nHeart system working\nBasic HUD visible'),

        ('Sprint 3\n(Polish & Ship)',
         'May 25–30',
         '• Full playtesting — all 5 endings reachable\n'
         '• Add sound effects + BGM (free assets ok)\n'
         '• Bug fixes, collision tuning, final art integration\n'
         '• Build web export, upload to submission link',
         'Final build exported\nAll endings tested\nGame submitted'),
    ]

    for sp in sprints:
        add_table_row(sp_tbl, sp)

    doc.add_paragraph()

    # ── 9. Git Workflow ───────────────────────────────────────────────────────
    heading(doc, '9. Git Workflow (Important!)', level=1)
    rules = [
        'Branch per feature: feat/minigame-professor, feat/minigame-niupai, feat/minigame-father, feat/art-assets',
        'Never push directly to main. Always open a Pull Request and ask Sirawee to review.',
        'Commit message format: feat: add professor circuit logic  OR  art: add boy walk sprites',
        'If you edit a .js file that Sirawee wrote, message him on LINE first.',
        'Art files only: artists can push directly to their feat/art-* branch without PR review.',
        'Pull from main every morning before starting work to avoid conflicts.',
    ]
    for r in rules:
        bullet(doc, r)

    # Save
    path = os.path.join(OUT_DIR, 'LOVEEXE_PMP_SRS.docx')
    doc.save(path)
    print(f'Saved: {path}')


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT 2 — Codebase Guide
# ══════════════════════════════════════════════════════════════════════════════

def make_codebase_guide():
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run('LOVE.EXE — Codebase Guide')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x5, 0x7B)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('How to clone, understand, and collaborate on this repository')
    r2.font.size = Pt(12)
    r2.italic = True

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run('Group 23  |  Written by Sirawee (110006226)  |  May 2026')
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ── 1. Quick Start ────────────────────────────────────────────────────────
    heading(doc, '1. Quick Start', level=1)
    body(doc, 'Step 1 — Clone the repository', bold=True)
    p = doc.add_paragraph()
    r = p.add_run('    git clone <repo-url>\n    cd LOVEexe')
    r.font.name = 'Courier New'
    r.font.size = Pt(10)

    body(doc, 'Step 2 — Open in Cocos Creator 2.4.8', bold=True)
    bullet(doc, 'Launch Cocos Creator 2.4.8 (must be this exact version — NOT 3.x)')
    bullet(doc, 'Click "Open Other Project" → select the LOVEexe folder')
    bullet(doc, 'Wait for the editor to import assets (~30 seconds first time)')

    body(doc, 'Step 3 — Open the MainMap scene', bold=True)
    bullet(doc, 'In the Assets panel: assets/scenes/MainMap.fire → double-click')
    bullet(doc, 'Press the Play button (▶) to test in browser')

    doc.add_paragraph()
    body(doc,
         'NOTE: Do NOT open this project in Cocos Creator 3.x — the JavaScript API is completely '
         'different and the project will not work.',
         bold=False, italic=True)

    # ── 2. Folder Structure ───────────────────────────────────────────────────
    heading(doc, '2. Folder Structure', level=1)

    tree = doc.add_paragraph()
    tree_run = tree.add_run(
        'LOVEexe/\n'
        '├── assets/\n'
        '│   ├── scripts/\n'
        '│   │   ├── core/           ← Framework scripts (DO NOT MODIFY without asking Sirawee)\n'
        '│   │   │   ├── GameManager.js\n'
        '│   │   │   ├── PlayerController.js\n'
        '│   │   │   ├── BoothTrigger.js\n'
        '│   │   │   └── EndingManager.js\n'
        '│   │   ├── minigames/      ← Each coder adds their mini-game script here\n'
        '│   │   │   ├── MiniGameBase.js   ← Base class — read this first!\n'
        '│   │   │   ├── MiniGame_Professor.js  (蘇昱銓 creates this)\n'
        '│   │   │   ├── MiniGame_NiuPai.js     (蘇昱銓 creates this)\n'
        '│   │   │   └── MiniGame_Father.js     (Anan creates this)\n'
        '│   │   └── ui/\n'
        '│   │       ├── DialogueSystem.js\n'
        '│   │       └── HUDController.js\n'
        '│   ├── resources/\n'
        '│   │   └── game-data.json  ← All dialogue text & ending slides live here\n'
        '│   ├── scenes/             ← .fire scene files (created in Cocos Editor)\n'
        '│   ├── textures/\n'
        '│   │   ├── characters/     ← 許夏綠 puts character PNGs here\n'
        '│   │   ├── map/            ← 張淑芬 puts tileset PNGs here\n'
        '│   │   ├── ui/             ← HUD icons, dialogue frames\n'
        '│   │   └── endings/        ← Ending CG images\n'
        '│   └── audio/\n'
        '│       ├── bgm/\n'
        '│       └── sfx/\n'
        '└── LOVEEXE_PMP_SRS.docx    ← This project\'s PMP+SRS document\n'
    )
    tree_run.font.name = 'Courier New'
    tree_run.font.size = Pt(9)

    # ── 3. Script Reference ───────────────────────────────────────────────────
    heading(doc, '3. Script Reference', level=1)

    scripts = [
        (
            'GameManager.js',
            'core/',
            'Singleton — attach to one empty Node in the FIRST scene (Title or MainMap). '
            'Automatically persists across all scene loads via cc.game.addPersistRootNode.\n\n'
            'Key properties:\n'
            '  • hearts (0–3) — Heart Fragments collected\n'
            '  • coins — total coins earned\n'
            '  • affinityScore (0–100) — determines ending branch\n'
            '  • miniGameState — tracks completion + best score per game\n'
            '  • easterEggs — boolean flags for special endings\n\n'
            'Key methods you will call:\n'
            '  • gm.completeMiniGame("professor", score)  → call this when player wins\n'
            '  • gm.loseLife()  → NOT in GameManager; call this.loseLife() from MiniGameBase\n'
            '  • gm.returnToMap()  → go back to campus after mini-game ends\n'
            '  • gm.goToEnding()  → call from final NPC after all mini-games done\n\n'
            'How to access from any script:\n'
            '  const gm = cc.find("GameManager").getComponent("GameManager");'
        ),
        (
            'MiniGameBase.js',
            'minigames/',
            'Abstract base class. DO NOT attach this directly to a node.\n\n'
            'Your mini-game script should start like this:\n\n'
            '  const MiniGameBase = require("MiniGameBase");\n'
            '  cc.Class({\n'
            '      extends: MiniGameBase,\n'
            '      properties: {\n'
            '          gameKey: { default: "professor" },  // MUST match GameManager key\n'
            '          gameDuration: { default: 60 },\n'
            '          maxLives: { default: 3 },\n'
            '          // ... your own properties\n'
            '      },\n'
            '      onGameStart() { /* spawn enemies, setup board */ },\n'
            '      onGameUpdate(dt) { /* per-frame logic */ },\n'
            '      onGameEnd(won) { /* cleanup */ },\n'
            '  });\n\n'
            'Inherited methods you can call from your script:\n'
            '  • this.addScore(points) — adds to score + updates label\n'
            '  • this.loseLife() — decrements lives, triggers GAME OVER at 0\n'
            '  • this.getScore() — returns current score\n\n'
            'The base class auto-calls GameManager.completeMiniGame() on win.'
        ),
        (
            'DialogueSystem.js',
            'ui/',
            'Attach to a persistent UI Canvas node. Wire up panel, labels, avatar sprite, '
            'choice container, and choiceButtonPrefab in the Inspector.\n\n'
            'To start a dialogue from any script:\n\n'
            '  const dlg = cc.find("Canvas/DialogueSystem").getComponent("DialogueSystem");\n'
            '  dlg.startDialogue("professor_pre", () => {\n'
            '      // callback: runs when dialogue is finished\n'
            '      GameManager.instance.loadMiniGame("professor");\n'
            '  });\n\n'
            'All dialogue content is in assets/resources/game-data.json — '
            'you can add new dialogue trees there without touching this script.'
        ),
        (
            'BoothTrigger.js',
            'core/',
            'Attach to a BoxCollider node on the campus map. Set gameKey to the correct enum value.\n\n'
            'When player walks into the collider → "Press E" prompt appears.\n'
            'When player presses E → loads the corresponding mini-game scene.\n\n'
            'The collision manager must be enabled:\n'
            '  cc.director.getCollisionManager().enabled = true;\n'
            '  cc.director.getCollisionManager().enabledDebugDraw = true;  // dev only'
        ),
        (
            'game-data.json',
            'resources/',
            'Single source of truth for ALL text content.\n\n'
            'Structure:\n'
            '  {\n'
            '    "dialogues": {\n'
            '      "tree_id": {\n'
            '        "lines": [\n'
            '          { "speaker": "Name", "text": "...", "avatar": "filename_no_ext",\n'
            '            "choices": [ { "text": "...", "affinityDelta": 10, "jumpTo": "other_tree" } ] }\n'
            '        ]\n'
            '      }\n'
            '    },\n'
            '    "endings": {\n'
            '      "ending_key": { "slides": [ { "title": "...", "body": "...", "image": "filename" } ] }\n'
            '    }\n'
            '  }\n\n'
            'To add new dialogue: just add a new tree to the "dialogues" object.\n'
            'Avatar images must exist at: assets/textures/characters/<avatar>.png'
        ),
    ]

    for name, folder, desc in scripts:
        heading(doc, f'{name}  [{folder}]', level=2)
        p = doc.add_paragraph(desc)
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.name = 'Courier New'
        doc.add_paragraph()

    # ── 4. How to Create Your Mini-game Scene ─────────────────────────────────
    heading(doc, '4. How to Create Your Mini-game Scene (Step by Step)', level=1)
    steps = [
        ('Step 1', 'In Cocos Creator: File → New Scene → name it "MiniGame_Professor" (or NiuPai / Father).'),
        ('Step 2', 'Add a Canvas node (should exist by default).'),
        ('Step 3', 'Create an empty Node named "GameLogic". Attach your MiniGame_XXX.js script to it.'),
        ('Step 4', 'In the Inspector, set gameKey to the correct string ("professor" / "niupai" / "father").'),
        ('Step 5', 'Create UI nodes: TimerLabel (Label), ScoreLabel (Label), LivesContainer (empty node with life icons). Drag them into the corresponding fields in your script\'s Inspector panel.'),
        ('Step 6', 'Create a ResultPanel node (hidden by default) with a title label, score label, and two buttons: "Retry" and "Return to Map". Wire button click events to onRetry() and onReturnToMap() on your GameLogic node.'),
        ('Step 7', 'Create a CountdownPanel with a label. This is shown "3-2-1" before the game starts.'),
        ('Step 8', 'Save the scene (Ctrl+S). The scene file (.fire) goes into assets/scenes/.'),
        ('Step 9', 'Test by calling GameManager.loadMiniGame("professor") from a button in MainMap, or temporarily change the first loaded scene in Build Settings.'),
    ]

    for s, d in steps:
        p = doc.add_paragraph()
        r1 = p.add_run(f'{s}: ')
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(d)
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # ── 5. Scene List ─────────────────────────────────────────────────────────
    heading(doc, '5. Scene List & Ownership', level=1)

    sc_tbl = doc.add_table(rows=1, cols=3)
    sc_tbl.style = 'Table Grid'
    table_header_row(sc_tbl, ['Scene File (.fire)', 'Description', 'Owner'])

    scenes = [
        ('Title.fire',               'Title screen with LOVE.EXE logo and "Start" button',    'Sirawee'),
        ('MainMap.fire',             'NTHU campus top-down map, player movement, 3 booths',    'Sirawee'),
        ('MiniGame_Professor.fire',  'Logic Gate Puzzle mini-game',                            '蘇昱銓'),
        ('MiniGame_NiuPai.fire',     'Grid Stealth / AI Chase mini-game',                      '蘇昱銓'),
        ('MiniGame_Father.fire',     'Rhythm / Catch mini-game',                               'Anan'),
        ('Ending.fire',              'Ending slides + credits roll',                           'Anan'),
    ]
    for s in scenes:
        add_table_row(sc_tbl, s)

    doc.add_paragraph()

    # ── 6. Common Mistakes ────────────────────────────────────────────────────
    heading(doc, '6. Common Mistakes to Avoid', level=1)

    mistakes = [
        ('Wrong Cocos version',
         'Using Cocos Creator 3.x — this project uses 2.4.8 JavaScript API. '
         'cc.director, cc.game, cc.loader, etc. are different in 3.x.'),
        ('Forgetting require()',
         'To use MiniGameBase in your script: const MiniGameBase = require("MiniGameBase"); '
         'at the top of your file. Same for any cross-script dependency.'),
        ('Hard-coding dialogue text',
         'Never put dialogue strings directly in your .js file. Add them to game-data.json and call '
         'DialogueSystem.startDialogue("your_tree_id", callback).'),
        ('Calling cc.director.loadScene() directly',
         'Always use GameManager methods: gm.loadMiniGame(), gm.returnToMap(), gm.goToEnding(). '
         'This ensures state is saved correctly before the scene change.'),
        ('Forgetting to set gameKey',
         'If your MiniGameBase subclass has gameKey set to the wrong value, the Heart Fragment '
         'will not be awarded because GameManager.miniGameState lookup will fail.'),
        ('Pushing art files to main directly',
         'Put PNGs on your feat/art-* branch. Art files can be large and clutter the main branch history.'),
        ('Scene files in wrong folder',
         'All .fire scene files must be inside assets/scenes/. Cocos Creator needs this path '
         'to find scenes by name in cc.director.loadScene("SceneName").'),
    ]

    for title, detail in mistakes:
        p = doc.add_paragraph()
        r1 = p.add_run(f'⚠  {title}: ')
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        r2 = p.add_run(detail)
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # ── 7. Contact ────────────────────────────────────────────────────────────
    heading(doc, '7. Questions?', level=1)
    body(doc, 'Contact Sirawee (110006226) on LINE or WeChat for:', size=10)
    bullet(doc, 'Git merge conflicts')
    bullet(doc, 'Questions about GameManager API')
    bullet(doc, 'Help wiring up Inspector fields')
    bullet(doc, 'Anything broken after pulling from main')

    path = os.path.join(OUT_DIR, 'LOVEEXE_Codebase_Guide.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    make_pmp_srs()
    make_codebase_guide()
    print('Both documents generated successfully.')
